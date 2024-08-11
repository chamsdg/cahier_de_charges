import streamlit as st
from io import BytesIO
from docx import Document
import pdfkit


# Fonction pour générer un fichier DOCX
def generate_docx(data):
    doc = Document()
    doc.add_heading('Cahier des Charges', 0)
    for key, value in data.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(str(value))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Fonction pour générer un fichier PDF
def generate_pdf(data):
    html_content = "<h1>Cahier des Charges</h1>"
    for key, value in data.items():
        html_content += f"<h2>{key}</h2><p>{str(value)}</p>"
    
    pdf = pdfkit.from_string(html_content, False)
    return pdf

# Interface utilisateur
# Interface utilisateur
def main():
    st.title("Cahier des Charges Formulaire")

    # Ajouter le texte d'introduction
    intro_text = """
    Ce document vous aidera à fournir à l'équipe Web de Techonologie IA les informations nécessaires pour planifier et développer efficacement votre nouveau site Web.
    L’équipe Web utilisera ces réponses de façon à ce que votre site Web réponde précisément à vos exigences.
    Ne remplissez que les informations concernant votre situation et ne vous inquiétez pas s'il reste des sections que vous n'êtes pas encore en mesure de remplir. Nous utiliserons ce document pour faciliter la discussion qui permettra de préciser le cahier des charges.
    Ce document deviendra notre contrat concernant ce à quoi ressemblera le site et sur sa date de livraison. Nous utilisons un processus réactif grâce auquel vous serez étroitement impliqué dans toutes les étapes du développement et qui vous permettra de voir le site se construire au fur et à mesure. Cette procédure limite les surprises éventuelles.
    """
    st.write(intro_text)
    
    with st.form("cahier_de_charges_form"):
        client_name = st.text_input("Nom du client")
        school_service = st.text_input("École/Service")
        function = st.text_input("Fonction")
        email = st.text_input("Adresse e-mail")
        phone_number = st.text_input("Numéro de téléphone")
        web_address = st.text_input("Adresse web")
        building = st.text_input("Bâtiment")
        project_description = st.text_area("Description du projet")
        target_audience = st.text_area("Public visé")
        essential_deliverables = st.text_area("Livrables essentiels")
        desired_deliverables = st.text_area("Livrables désirés")
        admin_tasks = st.text_area("Tâches administratives")
        admin_responsibility = st.text_area("Responsable administratif")
        skill_level = st.selectbox("Niveau de compétences en informatique", ["Débutant", "Intermédiaire", "Avancé"])
        site_map = st.radio("Avez-vous une carte du site ?", ["Oui", "Non", "Je ne sais pas"])
        section_estimate = st.number_input("Estimation du nombre de sections", min_value=0)
        page_estimate = st.number_input("Estimation du nombre de pages", min_value=0)
        style_guide = st.radio("Votre entreprise possède-t-elle un guide de style ?", ["Oui", "Non", "Je ne sais pas"])
        image_ideas = st.text_area("Idées d'images ou de couleurs")
        features = st.multiselect("Fonctionnalités souhaitées", ["Facile à mettre à jour", "Bien classé dans les recherches Google", "Optimisé pour téléphones mobiles", "Galeries photo et multimédia", "Formulaires de commentaires/contact", "Lettres d'information et abonnement", "Section réservée aux membres", "Vidéo/audio", "Calendrier", "Enquêtes", "Blog", "Autres"])
        accessibility = st.text_area("Accessibilité")
        content_types = st.text_area("Types de contenus")
        produced_content = st.text_area("Contenu produit jusqu'à présent")
        new_content = st.text_area("Nouveau contenu à produire")
        content_help = st.radio("Besoin d'aide pour produire le nouveau contenu ?", ["Oui", "Non", "Peut-être"])
        marketing_materials = st.text_area("Documents marketing associés")
        search_terms = st.text_area("Termes de recherche")
        social_media_strategy = st.radio("Possédez-vous une stratégie en matière de réseaux sociaux ?", ["Oui", "Non", "Je ne sais pas"])
        social_media_sites = st.text_area("Adresses des sites de réseaux sociaux")
        external_services = st.text_area("Services externes à intégrer")
        funding = st.radio("Votre site Web fait-il partie d'un projet de recherche financé ?", ["Oui", "Non"])
        funding_orgs = st.text_area("Organismes de financement")
        stakeholders = st.text_area("Parties prenantes")
        existing_developers = st.text_area("Développeurs Web existants")
        current_site = st.text_area("Satisfaction du site Web actuel")
        current_site_satisfaction = st.text_area("Points de satisfaction du site actuel")
        current_host = st.text_area("Hébergeur actuel")
        host_satisfaction = st.text_area("Satisfaction de l'hébergeur actuel")
        good_practices = st.text_area("Bonnes pratiques")
        liked_sites = st.text_area("Sites Web que vous aimez")
        other_comments = st.text_area("Autres commentaires")

        submit_button = st.form_submit_button("Soumettre")

    if submit_button:
        data = {
            "Nom du client": client_name,
            "École/Service": school_service,
            "Fonction": function,
            "Adresse e-mail": email,
            "Numéro de téléphone": phone_number,
            "Adresse web": web_address,
            "Bâtiment": building,
            "Description du projet": project_description,
            "Public visé": target_audience,
            "Livrables essentiels": essential_deliverables,
            "Livrables désirés": desired_deliverables,
            "Tâches administratives": admin_tasks,
            "Responsable administratif": admin_responsibility,
            "Niveau de compétences en informatique": skill_level,
            "Carte du site": site_map,
            "Estimation du nombre de sections": section_estimate,
            "Estimation du nombre de pages": page_estimate,
            "Guide de style": style_guide,
            "Idées d'images ou de couleurs": image_ideas,
            "Fonctionnalités souhaitées": ", ".join(features),
            "Accessibilité": accessibility,
            "Types de contenus": content_types,
            "Contenu produit jusqu'à présent": produced_content,
            "Nouveau contenu à produire": new_content,
            "Besoin d'aide pour produire le nouveau contenu": content_help,
            "Documents marketing associés": marketing_materials,
            "Termes de recherche": search_terms,
            "Stratégie en matière de réseaux sociaux": social_media_strategy,
            "Adresses des sites de réseaux sociaux": social_media_sites,
            "Services externes à intégrer": external_services,
            "Projet de recherche financé": funding,
            "Organismes de financement": funding_orgs,
            "Parties prenantes": stakeholders,
            "Développeurs Web existants": existing_developers,
            "Satisfaction du site Web actuel": current_site,
            "Points de satisfaction du site actuel": current_site_satisfaction,
            "Hébergeur actuel": current_host,
            "Satisfaction de l'hébergeur actuel": host_satisfaction,
            "Bonnes pratiques": good_practices,
            "Sites Web que vous aimez": liked_sites,
            "Autres commentaires": other_comments,
        }

        st.write("### Merci de telecharger:")
        docx_buffer = generate_docx(data)
        st.download_button(label="Télécharger en DOCX", data=docx_buffer, file_name="cahier_de_charges.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        #pdf_file = generate_pdf(data)
        #st.download_button(label="Télécharger en PDF", data=pdf_file, file_name="cahier_de_charges.pdf", mime="application/pdf")

if __name__ == "__main__":
    main()
